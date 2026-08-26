"""Docs Agent - creates Google Docs and Microsoft Word (.docx) documents."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional
from uuid import uuid4

import structlog

from ..db.database import Database
from ..db.repositories import DraftRepo
from ..llm.base import BaseLLM, Message
from ..quality import QualityCriterion, RefinementConfig, RefinementLoop
from ..quality.hooks import build_hook_chain, make_em_dash_hook, make_proofread_hook
from ..workspace.manager import WorkspaceManager
from .base import BaseAgent

logger = structlog.get_logger()

# Docs quality criteria
_DOCS_CRITERIA = [
    QualityCriterion(
        name="structure_quality",
        description="Logical structure, clear hierarchy, well-organized sections",
    ),
    QualityCriterion(
        name="brand_alignment", description="Consistent with brand voice and tone guidelines"
    ),
    QualityCriterion(
        name="clarity", description="Clear, concise writing free of jargon and ambiguity"
    ),
    QualityCriterion(
        name="completeness", description="All requested sections covered with substantive content"
    ),
    QualityCriterion(
        name="audience_appropriateness",
        description="Content tailored to the stated audience and purpose",
    ),
]


class DocsAgent(BaseAgent):
    """Creates Google Docs for marketing reports, briefs, proposals, and more.

    Uses the Google Docs API (via service account) to create live, editable
    documents and shares them via link.  Falls back gracefully when credentials
    are not configured.
    """

    agent_id = "docs_agent"
    agent_name = "Docs Agent"

    def __init__(
        self,
        llm: BaseLLM,
        db: Database,
        workspace_manager: WorkspaceManager,
        google_credentials_path: str = "",
        google_oauth_token_path: str = "",
        media_storage: Optional[Any] = None,
        proofreader: Optional[Any] = None,
        scanning_llm: Optional[BaseLLM] = None,
        max_iterations: int = 10,
    ) -> None:
        self._workspace_mgr = workspace_manager
        self._google_credentials_path = google_credentials_path
        self._google_oauth_token_path = google_oauth_token_path
        self._media_storage = media_storage
        self._proofreader = proofreader
        self._scanning_llm = scanning_llm
        self._docs_service: Optional[Any] = None
        self._drive_service: Optional[Any] = None
        self._drafts = DraftRepo(db)
        super().__init__(llm=llm, db=db, max_iterations=max_iterations)

    # ── Google API lazy init ──────────────────────────────────────────────

    def _get_google_services(self) -> bool:
        """Lazily initialize Google Docs and Drive API services.

        Returns True if services are available, False otherwise.
        Uses OAuth2 token (preferred) or service account credentials (fallback).
        """
        if self._docs_service is not None:
            return True

        try:
            from googleapiclient.discovery import build

            from ..google_auth import get_google_credentials

            SCOPES = [
                "https://www.googleapis.com/auth/documents",
                "https://www.googleapis.com/auth/drive.file",
            ]
            # Force OAuth — service account lacks Docs API permission.
            # All working scripts (create_monetization_doc, create_entity_setup_playbook, etc.)
            # use service_account_path="" to force OAuth for Docs API access.
            credentials = get_google_credentials(
                oauth_token_path=self._google_oauth_token_path,
                service_account_path="",
                scopes=SCOPES,
            )
            if credentials is None:
                return False

            self._docs_service = build("docs", "v1", credentials=credentials)
            self._drive_service = build("drive", "v3", credentials=credentials)
            logger.info("docs_google_services_initialized")
            return True
        except Exception as e:
            logger.error("docs_google_services_init_failed", error=str(e))
            return False

    def _share_document(self, doc_id: str) -> None:
        """Share a document with authorized users only (restricted access)."""
        from ..google_auth import share_file_restricted

        share_file_restricted(self._drive_service, doc_id)

    # ── Output format dispatch ────────────────────────────────────────────

    def _resolve_format(self, requested: str) -> str:
        """Determine output format. Default to google_docs when credentials exist.

        Accepted values:
            ""              auto: google_docs if creds, otherwise docx
            "google_docs"   Google Docs only
            "docx"          Microsoft Word only
            "both"          Build both formats
        """
        if requested:
            r = requested.lower().strip()
            if r in ("google", "gdoc", "gdocs"):
                return "google_docs"
            if r in ("word", "office", "ms_word", "microsoft_word"):
                return "docx"
            return r
        if self._google_credentials_path and Path(self._google_credentials_path).exists():
            return "google_docs"
        return "docx"

    async def _build_and_export(
        self,
        structure: Dict[str, Any],
        workspace_id: str,
        output_format: str,
    ) -> Dict[str, Any]:
        """Build the doc and export in the requested format(s)."""
        if output_format == "google_docs":
            return self._build_google_doc(structure, workspace_id)

        if output_format == "docx":
            return await self._build_docx(structure, workspace_id)

        if output_format == "both":
            gdoc_result = self._build_google_doc(structure, workspace_id)
            docx_result = await self._build_docx(structure, workspace_id)
            merged: Dict[str, Any] = {"status": "created"}
            if gdoc_result.get("google_docs_url"):
                merged["google_docs_url"] = gdoc_result["google_docs_url"]
                merged["google_docs_id"] = gdoc_result.get("google_docs_id")
            if docx_result.get("file_path"):
                merged["file_path"] = docx_result["file_path"]
            if docx_result.get("cdn_url"):
                merged["cdn_url"] = docx_result["cdn_url"]
            if gdoc_result.get("status") != "created" and docx_result.get("status") != "created":
                merged["status"] = "error"
                merged["message"] = "Both Google Docs and .docx generation failed"
            return merged

        return {
            "status": "error",
            "message": f"Unknown output_format: {output_format!r}",
        }

    # ── System prompt ─────────────────────────────────────────────────────

    def get_system_prompt(self, workspace_id: Optional[str] = None) -> str:
        ws_label = workspace_id or "all workspaces"
        gdocs_available = bool(self._google_credentials_path)
        return f"""You are the Docs Agent for {ws_label}.
You create professional Google Docs for marketing teams.

Google Docs API: {"configured" if gdocs_available else "not configured"}

CRITICAL RULES:
- You MUST ALWAYS call the create_document or create_document_from_content tool. NEVER generate document content as text in your response.
- Your ONLY job is to call a tool that creates a real Google Doc and return the link. Do NOT write out document content yourself.
- If the tool returns "not_configured", report that Google credentials need to be set up. Do NOT fall back to generating content as text.
- After creating a document, your response should be brief: state that the document was created and include the Google Docs URL.

You can create:
- Marketing briefs and strategy documents
- Reports and analyses
- Proposals and pitch documents
- Blog posts and long-form content
- Newsletters and email copy
- Meeting notes and agendas

When creating documents:
- Use clear section headings and hierarchy
- Keep paragraphs concise and scannable
- Include executive summaries for long documents
- Inject brand voice when available
- Use professional formatting (headings, subheadings, bullet points)
- ABBREVIATIONS: Always spell out abbreviations/acronyms on first use with the abbreviation in parentheses — e.g., "Total Addressable Market (TAM)". After the first occurrence, use the abbreviation only.
- ANTI-AI WRITING (MANDATORY):
  • NEVER use em dashes (—) or en dashes (–). Use commas, periods, semicolons, or " - " (spaced hyphen) instead.
  • NEVER bold words in the middle of a sentence. Bolding is only for section headers or standalone labels.

Output is a live Google Doc link that opens directly in the browser.
Current workspace: {ws_label}"""

    # ── Tools ─────────────────────────────────────────────────────────────

    def register_tools(self) -> None:

        @self.tool_registry.register(
            name="create_document",
            description=(
                "Create a document from a topic. Generates professional document structure "
                "with headings, sections, and brand-appropriate content. "
                "Defaults to Google Docs when configured, otherwise .docx. "
                "Set output_format to 'google_docs', 'docx', or 'both'."
            ),
        )
        async def create_document(
            topic: str,
            workspace_id: str,
            purpose: str = "report",
            audience: str = "general",
            sections: Optional[List[str]] = None,
            output_format: str = "",
        ) -> Dict[str, Any]:
            fmt = self._resolve_format(output_format)
            brand_voice = await self._workspace_mgr.get_brand_voice(workspace_id)
            reference_docs = await self._workspace_mgr.get_reference_docs(workspace_id)

            structure = await self._generate_doc_structure(
                topic=topic,
                purpose=purpose,
                audience=audience,
                brand_voice=brand_voice,
                reference_docs=reference_docs,
                sections=sections,
            )

            if self._proofreader:
                try:
                    structure, _ = await self._proofreader.proofread_dict_fields(
                        structure,
                        fields=["title", "heading", "body"],
                    )
                except Exception as e:
                    logger.warning("docs_proofread_failed", error=str(e))

            result = await self._build_and_export(structure, workspace_id, fmt)
            return {
                **result,
                "topic": topic,
                "purpose": purpose,
                "workspace_id": workspace_id,
                "output_format": fmt,
            }

        @self.tool_registry.register(
            name="create_document_from_content",
            description=(
                "Create a document from existing text content. Structures the content "
                "into a well-formatted document with headings and sections. "
                "Defaults to Google Docs when configured, otherwise .docx. "
                "Set output_format to 'google_docs', 'docx', or 'both'."
            ),
        )
        async def create_document_from_content(
            content: str,
            workspace_id: str,
            title: str = "",
            format_style: str = "professional",
            output_format: str = "",
        ) -> Dict[str, Any]:
            fmt = self._resolve_format(output_format)
            brand_voice = await self._workspace_mgr.get_brand_voice(workspace_id)

            prompt = f"""Convert the following content into a structured document.

CONTENT:
{content[:4000]}

TITLE: {title or "(generate an appropriate title)"}
STYLE: {format_style}

BRAND CONTEXT:
{brand_voice[:300] if brand_voice else "No brand voice specified."}

{_doc_structure_format_instructions()}"""

            response = await self.llm.complete(
                messages=[Message(role="user", content=prompt)],
                temperature=0.7,
            )

            structure = _parse_doc_structure(response.content)

            if self._proofreader:
                try:
                    structure, _ = await self._proofreader.proofread_dict_fields(
                        structure,
                        fields=["title", "heading", "body"],
                    )
                except Exception as e:
                    logger.warning("docs_proofread_failed", error=str(e))

            result = await self._build_and_export(structure, workspace_id, fmt)
            return {
                **result,
                "title": structure.get("title", title),
                "workspace_id": workspace_id,
                "output_format": fmt,
            }

        @self.tool_registry.register(
            name="export_draft_to_doc",
            description=(
                "Export an existing draft from the database to a document. "
                "Fetches the draft content and creates a formatted document. "
                "Defaults to Google Docs when configured, otherwise .docx. "
                "Set output_format to 'google_docs', 'docx', or 'both'."
            ),
        )
        async def export_draft_to_doc(
            draft_id: str,
            workspace_id: str,
            output_format: str = "",
        ) -> Dict[str, Any]:
            fmt = self._resolve_format(output_format)
            draft = await self._drafts.get(draft_id)
            if not draft:
                return {"status": "error", "message": f"Draft '{draft_id}' not found"}

            structure = {
                "title": draft.title or f"Draft: {draft_id}",
                "sections": [{"heading": draft.title or "Content", "body": draft.body, "level": 1}],
            }

            if self._proofreader:
                try:
                    structure, _ = await self._proofreader.proofread_dict_fields(
                        structure,
                        fields=["title", "heading", "body"],
                    )
                except Exception as e:
                    logger.warning("docs_proofread_failed", error=str(e))

            result = await self._build_and_export(structure, workspace_id, fmt)
            return {
                **result,
                "draft_id": draft_id,
                "workspace_id": workspace_id,
                "output_format": fmt,
            }

    # ── Internal methods ──────────────────────────────────────────────────

    async def _generate_doc_structure(
        self,
        topic: str,
        purpose: str,
        audience: str,
        brand_voice: str,
        reference_docs: str = "",
        sections: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        """Use the LLM to generate a structured JSON document outline."""
        sections_hint = ""
        if sections:
            sections_hint = f"\nREQUESTED SECTIONS: {', '.join(sections)}"

        ref_docs_section = ""
        if reference_docs:
            ref_docs_section = f"""
REFERENCE DOCUMENTS (use real data, metrics, and facts from these):
{reference_docs[:15000]}
"""

        prompt = f"""Create a professional document structure.

TOPIC: {topic}
PURPOSE: {purpose}
AUDIENCE: {audience}{sections_hint}

BRAND CONTEXT:
{brand_voice[:400] if brand_voice else "No brand voice specified."}
{ref_docs_section}

{_doc_structure_format_instructions()}"""

        config = RefinementConfig(
            criteria=_DOCS_CRITERIA,
            max_iterations=2,
            quality_threshold=7.0,
            log_prefix="docs_refinement",
        )
        hooks = build_hook_chain(
            make_em_dash_hook(),
            make_proofread_hook(self._proofreader, "document content")
            if self._proofreader
            else None,
        )
        loop = RefinementLoop(
            config=config,
            generation_llm=self.llm,
            critique_llm=self._scanning_llm,
        )
        result = await loop.generate_and_refine(prompt, brand_voice, hooks)

        return _parse_doc_structure(result.content)

    def _build_google_doc(self, structure: Dict[str, Any], workspace_id: str) -> Dict[str, Any]:
        """Create a Google Doc via the Docs API and return its URL."""
        if not self._get_google_services():
            return {
                "status": "not_configured",
                "message": (
                    "Google credentials not configured. "
                    "Set GOOGLE_CREDENTIALS_PATH in .env to a service account JSON file."
                ),
            }

        title = structure.get("title", "Untitled Document")
        include_toc = structure.get("include_toc", False)
        sections = _normalize_sections(structure, skip_toc=include_toc)

        try:
            # 1. Create empty document
            doc = self._docs_service.documents().create(body={"title": title}).execute()
            doc_id = doc["documentId"]

            # 2. Build requests using a sequential cursor
            requests: List[Dict[str, Any]] = []
            cursor = 1  # Docs API index starts at 1

            for section in sections:
                heading = section.get("heading", "")
                level = section.get("level", 1)
                content_blocks = section.get("content", [])

                # Insert heading
                if heading:
                    heading_text = heading + "\n"
                    requests.append(
                        {"insertText": {"location": {"index": cursor}, "text": heading_text}}
                    )
                    h_start = cursor
                    h_end = cursor + len(heading_text)
                    requests.append(
                        {
                            "updateParagraphStyle": {
                                "range": {"startIndex": h_start, "endIndex": h_end},
                                "paragraphStyle": {
                                    "namedStyleType": _heading_style(level),
                                    "spaceAbove": {"magnitude": 12, "unit": "PT"},
                                    "spaceBelow": {"magnitude": 4, "unit": "PT"},
                                },
                                "fields": "namedStyleType,spaceAbove,spaceBelow",
                            }
                        }
                    )
                    cursor = h_end

                # Insert each content block
                for block in content_blocks:
                    block_type = block.get("type", "paragraph")

                    if block_type == "paragraph":
                        text = block.get("text", "")
                        if not text:
                            continue
                        para_text = text + "\n"
                        plain, formatting_runs = _parse_inline_formatting(para_text)
                        requests.append(
                            {"insertText": {"location": {"index": cursor}, "text": plain}}
                        )
                        p_start = cursor
                        p_end = cursor + len(plain)
                        requests.append(
                            {
                                "updateParagraphStyle": {
                                    "range": {"startIndex": p_start, "endIndex": p_end},
                                    "paragraphStyle": {
                                        "namedStyleType": "NORMAL_TEXT",
                                        "spaceBelow": {"magnitude": 6, "unit": "PT"},
                                    },
                                    "fields": "namedStyleType,spaceBelow",
                                }
                            }
                        )
                        # Apply inline bold/italic
                        for run_start, run_end, style in formatting_runs:
                            abs_start = cursor + run_start
                            abs_end = cursor + run_end
                            if style == "bold":
                                requests.append(
                                    {
                                        "updateTextStyle": {
                                            "range": {
                                                "startIndex": abs_start,
                                                "endIndex": abs_end,
                                            },
                                            "textStyle": {"bold": True},
                                            "fields": "bold",
                                        }
                                    }
                                )
                            elif style == "italic":
                                requests.append(
                                    {
                                        "updateTextStyle": {
                                            "range": {
                                                "startIndex": abs_start,
                                                "endIndex": abs_end,
                                            },
                                            "textStyle": {"italic": True},
                                            "fields": "italic",
                                        }
                                    }
                                )
                        cursor = p_end

                    elif block_type in ("bullets", "numbered_list"):
                        items = block.get("items", [])
                        if not items:
                            continue
                        list_start = cursor
                        for item in items:
                            item_text = str(item) + "\n"
                            requests.append(
                                {
                                    "insertText": {
                                        "location": {"index": cursor},
                                        "text": item_text,
                                    }
                                }
                            )
                            cursor += len(item_text)
                        list_end = cursor
                        # Apply bullet/numbered list style
                        preset = (
                            "BULLET_DISC_CIRCLE_SQUARE"
                            if block_type == "bullets"
                            else "NUMBERED_DECIMAL_NESTED"
                        )
                        requests.append(
                            {
                                "createParagraphBullets": {
                                    "range": {
                                        "startIndex": list_start,
                                        "endIndex": list_end,
                                    },
                                    "bulletPreset": preset,
                                }
                            }
                        )

                    elif block_type == "table":
                        t_headers = block.get("headers", [])
                        t_rows = block.get("rows", [])
                        if not t_headers:
                            continue
                        num_cols = len(t_headers)
                        num_rows = 1 + len(t_rows)  # header row + data rows
                        requests.append(
                            {
                                "insertTable": {
                                    "rows": num_rows,
                                    "columns": num_cols,
                                    "location": {"index": cursor},
                                }
                            }
                        )
                        # After insertTable, the table occupies document space.
                        # We need to populate cells. Each table cell starts at a
                        # known offset. The table structure:
                        #   table_start_idx = cursor + 1 (table element)
                        #   row 0: cell 0 starts at cursor + 4
                        # Rather than computing exact indices (fragile), we'll
                        # execute what we have so far, re-read the doc to get
                        # table indices, then populate.
                        # For simplicity, batch execute here and rebuild cursor.
                        if requests:
                            self._docs_service.documents().batchUpdate(
                                documentId=doc_id, body={"requests": requests}
                            ).execute()
                            requests = []

                        # Re-read doc to find the table we just inserted
                        doc_state = self._docs_service.documents().get(documentId=doc_id).execute()
                        table_element = None
                        for elem in doc_state.get("body", {}).get("content", []):
                            if "table" in elem:
                                table_element = elem
                        if table_element:
                            table = table_element["table"]
                            all_table_rows = [t_headers] + t_rows
                            # Collect (cell_index, cell_text) pairs, then
                            # insert in REVERSE index order so earlier inserts
                            # don't shift later cell positions.
                            cell_inserts: List[tuple] = []
                            for ri, row_data in enumerate(all_table_rows):
                                if ri >= len(table.get("tableRows", [])):
                                    break
                                table_row = table["tableRows"][ri]
                                for ci, cell_val in enumerate(row_data):
                                    if ci >= len(table_row.get("tableCells", [])):
                                        break
                                    cell = table_row["tableCells"][ci]
                                    cell_content = cell.get("content", [])
                                    if cell_content:
                                        cell_idx = cell_content[0].get("startIndex", 0)
                                        cell_text = str(cell_val)
                                        if cell_text:
                                            cell_inserts.append((cell_idx, cell_text))
                            # Sort descending by index so inserts don't shift
                            cell_inserts.sort(key=lambda x: x[0], reverse=True)
                            for cell_idx, cell_text in cell_inserts:
                                requests.append(
                                    {
                                        "insertText": {
                                            "location": {"index": cell_idx},
                                            "text": cell_text,
                                        }
                                    }
                                )
                            # Execute text inserts
                            if requests:
                                self._docs_service.documents().batchUpdate(
                                    documentId=doc_id, body={"requests": requests}
                                ).execute()
                                requests = []

                            # Re-read doc to apply bold to header row cells
                            doc_state2 = (
                                self._docs_service.documents().get(documentId=doc_id).execute()
                            )
                            table_elem2 = None
                            for elem in doc_state2.get("body", {}).get("content", []):
                                if "table" in elem:
                                    table_elem2 = elem
                            if table_elem2 and table_elem2["table"].get("tableRows"):
                                header_row = table_elem2["table"]["tableRows"][0]
                                for hcell in header_row.get("tableCells", []):
                                    for para in hcell.get("content", []):
                                        si = para.get("startIndex", 0)
                                        ei = para.get("endIndex", si)
                                        if ei > si:
                                            requests.append(
                                                {
                                                    "updateTextStyle": {
                                                        "range": {
                                                            "startIndex": si,
                                                            "endIndex": ei,
                                                        },
                                                        "textStyle": {"bold": True},
                                                        "fields": "bold",
                                                    }
                                                }
                                            )
                            if requests:
                                self._docs_service.documents().batchUpdate(
                                    documentId=doc_id, body={"requests": requests}
                                ).execute()
                                requests = []

                        # Re-read doc to get updated cursor position
                        doc_state = self._docs_service.documents().get(documentId=doc_id).execute()
                        body_content = doc_state.get("body", {}).get("content", [])
                        if body_content:
                            last_elem = body_content[-1]
                            cursor = last_elem.get("endIndex", cursor) - 1
                        # Add a newline after the table
                        requests.append(
                            {"insertText": {"location": {"index": cursor}, "text": "\n"}}
                        )
                        cursor += 1

            # 3. Execute remaining batch requests
            if requests:
                self._docs_service.documents().batchUpdate(
                    documentId=doc_id, body={"requests": requests}
                ).execute()

            # 3b. Insert table of contents if requested
            if include_toc:
                self._insert_table_of_contents(doc_id)

            # 4. Share document
            self._share_document(doc_id)

            # 5. Move into CMO Agent folder
            from ..google_auth import ensure_drive_folder, move_file_to_folder

            folder_id = ensure_drive_folder(self._drive_service)
            if folder_id:
                move_file_to_folder(self._drive_service, doc_id, folder_id)

            docs_url = f"https://docs.google.com/document/d/{doc_id}/edit"
            logger.info(
                "google_doc_created",
                doc_id=doc_id,
                sections=len(sections),
                workspace=workspace_id,
            )

            return {
                "status": "created",
                "google_docs_url": docs_url,
                "google_docs_id": doc_id,
            }

        except Exception as e:
            logger.error("google_doc_build_failed", error=str(e))
            return {"status": "error", "message": str(e)}

    async def _build_docx(
        self, structure: Dict[str, Any], workspace_id: str
    ) -> Dict[str, Any]:
        """Build a Microsoft Word (.docx) file mirroring the doc structure.

        Mirrors the same structure consumed by _build_google_doc:
            { title, sections: [{ heading, level, content: [{type, ...}] }] }
        Supports paragraph (with **bold**/*italic*), bullets, numbered_list, table.
        Saves to data/docs/ and optionally uploads to the configured media store.
        """
        from docx import Document
        from docx.shared import Pt

        title = structure.get("title", "Untitled Document")
        include_toc = structure.get("include_toc", False)
        sections = _normalize_sections(structure, skip_toc=include_toc)

        try:
            document = Document()

            # Title
            title_para = document.add_heading(title, level=0)
            for run in title_para.runs:
                run.font.size = Pt(24)

            # Optional table of contents placeholder
            if include_toc:
                document.add_paragraph(
                    "Table of Contents",
                    style=document.styles["Heading 1"],
                )
                document.add_paragraph(
                    "(In Word, go to References > Table of Contents to insert "
                    "an auto-generated TOC. The headings below are styled to "
                    "be picked up automatically.)"
                ).italic = True

            for section in sections:
                heading = section.get("heading", "")
                level = section.get("level", 1)
                content_blocks = section.get("content", [])

                if heading:
                    # python-docx supports heading levels 0-9
                    h_level = max(1, min(int(level) if isinstance(level, int) else 1, 9))
                    document.add_heading(heading, level=h_level)

                for block in content_blocks:
                    block_type = block.get("type", "paragraph")

                    if block_type == "paragraph":
                        text = block.get("text", "")
                        if not text:
                            continue
                        plain, formatting_runs = _parse_inline_formatting(text)
                        # Strip trailing newline since add_paragraph already handles it
                        if plain.endswith("\n"):
                            plain = plain[:-1]
                        para = document.add_paragraph()
                        if formatting_runs:
                            cursor = 0
                            # Build a sorted list of (start, end, style) and emit runs
                            runs_sorted = sorted(formatting_runs, key=lambda r: r[0])
                            for r_start, r_end, style in runs_sorted:
                                if r_start > cursor:
                                    para.add_run(plain[cursor:r_start])
                                run = para.add_run(plain[r_start:r_end])
                                if style == "bold":
                                    run.bold = True
                                elif style == "italic":
                                    run.italic = True
                                cursor = r_end
                            if cursor < len(plain):
                                para.add_run(plain[cursor:])
                        else:
                            para.add_run(plain)

                    elif block_type in ("bullets", "numbered_list"):
                        items = block.get("items", [])
                        style_name = (
                            "List Bullet" if block_type == "bullets" else "List Number"
                        )
                        for item in items:
                            document.add_paragraph(str(item), style=style_name)

                    elif block_type == "table":
                        headers = block.get("headers", [])
                        rows = block.get("rows", [])
                        if not headers and not rows:
                            continue
                        cols = max(len(headers), max((len(r) for r in rows), default=0))
                        if cols == 0:
                            continue
                        table = document.add_table(rows=1, cols=cols)
                        table.style = "Light Grid Accent 1"
                        # Header row
                        hdr_cells = table.rows[0].cells
                        for ci in range(cols):
                            hdr_cells[ci].text = str(headers[ci]) if ci < len(headers) else ""
                            for paragraph in hdr_cells[ci].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                        # Data rows
                        for row_data in rows:
                            row_cells = table.add_row().cells
                            for ci in range(cols):
                                row_cells[ci].text = (
                                    str(row_data[ci]) if ci < len(row_data) else ""
                                )

                    elif block_type == "page_break":
                        document.add_page_break()

                    elif block_type == "heading":
                        sub_heading = block.get("text", "")
                        sub_level = block.get("level", 2)
                        if sub_heading:
                            document.add_heading(sub_heading, level=max(1, min(sub_level, 9)))

            # Save
            docs_dir = Path("./data/docs")
            docs_dir.mkdir(parents=True, exist_ok=True)
            slug = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-")[:40] or "document"
            filename = f"{workspace_id}_{slug}_{uuid4().hex[:6]}.docx"
            file_path = docs_dir / filename
            document.save(str(file_path))

            logger.info(
                "docx_created",
                file_path=str(file_path),
                sections=len(sections),
                workspace=workspace_id,
            )

            result: Dict[str, Any] = {
                "status": "created",
                "file_path": str(file_path),
            }

            if self._media_storage:
                try:
                    asset = await self._media_storage.upload_from_file(
                        file_path=str(file_path),
                        workspace_id=workspace_id,
                        media_type="doc",
                        source_agent=self.agent_id,
                        prompt=title,
                    )
                    if asset and asset.cdn_url:
                        result["cdn_url"] = asset.cdn_url
                except Exception as e:
                    logger.warning("docx_cdn_upload_failed", error=str(e))

            return result

        except Exception as e:
            logger.error("docx_build_failed", error=str(e))
            return {"status": "error", "message": str(e)}

    def _insert_table_of_contents(
        self,
        doc_id: str,
        batch_fn: Optional[Any] = None,
    ) -> None:
        """Insert a formatted table of contents at the beginning of a Google Doc.

        Reads all headings (HEADING_1 through HEADING_3) from the document,
        generates a TOC with internal heading links styled like a standard
        document TOC (11pt body text, indented sub-headings, linked entries).

        Note: The Google Docs API does not support inserting a native TOC
        element (no ``insertTableOfContents`` request exists — Google Issue
        #36758222, status: blocked). Page numbers are also unavailable since
        pagination is resolved only at render time. This method builds a
        professional linked TOC and adds an instruction for the user to
        replace it with a native TOC via Insert > Table of Contents if they
        need page numbers.

        Args:
            doc_id: The Google Doc document ID.
            batch_fn: Optional callable for rate-limited batchUpdate. If None,
                uses self._docs_service directly.
        """

        def _batch(reqs: List[Dict[str, Any]]) -> Any:
            if batch_fn:
                return batch_fn(reqs)
            return (
                self._docs_service.documents()
                .batchUpdate(documentId=doc_id, body={"requests": reqs})
                .execute()
            )

        # Read document to find all headings
        doc = self._docs_service.documents().get(documentId=doc_id).execute()
        body_content = doc.get("body", {}).get("content", [])

        headings: List[Dict[str, Any]] = []
        for elem in body_content:
            if "paragraph" not in elem:
                continue
            para = elem["paragraph"]
            style = para.get("paragraphStyle", {})
            named_style = style.get("namedStyleType", "")
            heading_id = style.get("headingId", "")
            if not named_style.startswith("HEADING_") or not heading_id:
                continue
            level = int(named_style.replace("HEADING_", ""))
            if level > 3:
                continue  # Only include H1-H3 in TOC
            text = ""
            for el in para.get("elements", []):
                text += el.get("textRun", {}).get("content", "")
            text = text.strip()
            if text:
                headings.append({"text": text, "level": level, "heading_id": heading_id})

        if not headings:
            return

        # -- Build TOC text ------------------------------------------------
        # Title line + tip line + entry lines + separator
        toc_title = "Table of Contents\n"
        toc_tip = (
            "For page numbers: select this TOC, delete it, then "
            "Insert \u2192 Table of contents \u2192 With page numbers.\n"
        )
        entry_lines: List[str] = []
        for h in headings:
            entry_lines.append(h["text"] + "\n")
        separator = "\n"  # blank line between TOC and body

        toc_text = toc_title + toc_tip + "".join(entry_lines) + separator

        # -- Insert all TOC text at index 1 --------------------------------
        requests: List[Dict[str, Any]] = [
            {"insertText": {"location": {"index": 1}, "text": toc_text}},
        ]

        # -- Style title: bold 16pt, NORMAL_TEXT (not a heading) -----------
        title_start = 1
        title_end = title_start + len(toc_title)
        requests.append(
            {
                "updateParagraphStyle": {
                    "range": {"startIndex": title_start, "endIndex": title_end},
                    "paragraphStyle": {
                        "namedStyleType": "NORMAL_TEXT",
                        "spaceBelow": {"magnitude": 4, "unit": "PT"},
                    },
                    "fields": "namedStyleType,spaceBelow",
                }
            }
        )
        requests.append(
            {
                "updateTextStyle": {
                    "range": {"startIndex": title_start, "endIndex": title_end - 1},
                    "textStyle": {
                        "bold": True,
                        "fontSize": {"magnitude": 16, "unit": "PT"},
                    },
                    "fields": "bold,fontSize",
                }
            }
        )

        # -- Style tip line: italic 9pt grey --------------------------------
        tip_start = title_end
        tip_end = tip_start + len(toc_tip)
        requests.append(
            {
                "updateParagraphStyle": {
                    "range": {"startIndex": tip_start, "endIndex": tip_end},
                    "paragraphStyle": {
                        "namedStyleType": "NORMAL_TEXT",
                        "spaceBelow": {"magnitude": 6, "unit": "PT"},
                    },
                    "fields": "namedStyleType,spaceBelow",
                }
            }
        )
        requests.append(
            {
                "updateTextStyle": {
                    "range": {"startIndex": tip_start, "endIndex": tip_end - 1},
                    "textStyle": {
                        "italic": True,
                        "fontSize": {"magnitude": 9, "unit": "PT"},
                        "foregroundColor": {
                            "color": {"rgbColor": {"red": 0.6, "green": 0.6, "blue": 0.6}}
                        },
                    },
                    "fields": "italic,fontSize,foregroundColor",
                }
            }
        )

        # -- Style each entry: 11pt, linked, indented by level -------------
        cursor = tip_end
        for h in headings:
            entry_text = h["text"] + "\n"
            entry_start = cursor
            entry_end = cursor + len(entry_text) - 1  # exclude newline for link

            if entry_end > entry_start:
                # Set paragraph to NORMAL_TEXT with tight spacing
                requests.append(
                    {
                        "updateParagraphStyle": {
                            "range": {
                                "startIndex": entry_start,
                                "endIndex": entry_start + len(entry_text),
                            },
                            "paragraphStyle": {
                                "namedStyleType": "NORMAL_TEXT",
                                "spaceAbove": {"magnitude": 2, "unit": "PT"},
                                "spaceBelow": {"magnitude": 2, "unit": "PT"},
                                "indentStart": {
                                    "magnitude": 18 * (h["level"] - 1),
                                    "unit": "PT",
                                },
                            },
                            "fields": "namedStyleType,spaceAbove,spaceBelow,indentStart",
                        }
                    }
                )

                # Set font to 11pt (standard body size)
                style_fields = "fontSize,link"
                text_style: Dict[str, Any] = {
                    "fontSize": {"magnitude": 11, "unit": "PT"},
                    "link": {"headingId": h["heading_id"]},
                }
                # H1 entries are bold
                if h["level"] == 1:
                    text_style["bold"] = True
                    style_fields += ",bold"

                requests.append(
                    {
                        "updateTextStyle": {
                            "range": {
                                "startIndex": entry_start,
                                "endIndex": entry_end,
                            },
                            "textStyle": text_style,
                            "fields": style_fields,
                        }
                    }
                )

            cursor += len(entry_text)

        _batch(requests)
        logger.info("google_doc_toc_inserted", doc_id=doc_id, entries=len(headings))

    def _update_google_doc(
        self, doc_id: str, structure: Dict[str, Any], workspace_id: str
    ) -> Dict[str, Any]:
        """Update an existing Google Doc by clearing its content and rebuilding."""
        import time

        if not self._get_google_services():
            return {
                "status": "not_configured",
                "message": "Google credentials not configured.",
            }

        # Track write ops for rate limiting (Google Docs: 60 writes/min/user)
        _write_timestamps: List[float] = []

        def _rate_limited_batch(requests_payload: List[Dict[str, Any]]) -> Any:
            """Execute batchUpdate with rate limiting."""
            now = time.time()
            # Remove timestamps older than 60 seconds
            while _write_timestamps and now - _write_timestamps[0] > 60:
                _write_timestamps.pop(0)
            # If we've hit 55 requests in the last 60s, sleep until the window clears
            if len(_write_timestamps) >= 55:
                sleep_time = 61 - (now - _write_timestamps[0])
                if sleep_time > 0:
                    logger.info("google_docs_rate_limit_pause", sleep_seconds=round(sleep_time, 1))
                    time.sleep(sleep_time)
            result = (
                self._docs_service.documents()
                .batchUpdate(documentId=doc_id, body={"requests": requests_payload})
                .execute()
            )
            _write_timestamps.append(time.time())
            return result

        try:
            # Read existing doc to find content length
            doc = self._docs_service.documents().get(documentId=doc_id).execute()
            body_content = doc.get("body", {}).get("content", [])
            if body_content:
                end_index = body_content[-1].get("endIndex", 1)
            else:
                end_index = 1

            # Delete all existing content (index 1 to end-1)
            if end_index > 2:
                _rate_limited_batch(
                    [
                        {
                            "deleteContentRange": {
                                "range": {
                                    "startIndex": 1,
                                    "endIndex": end_index - 1,
                                }
                            }
                        }
                    ]
                )

            include_toc = structure.get("include_toc", False)
            sections = _normalize_sections(structure, skip_toc=include_toc)

            # Rebuild content using same logic as _build_google_doc
            requests: List[Dict[str, Any]] = []
            cursor = 1

            for section in sections:
                heading = section.get("heading", "")
                level = section.get("level", 1)
                content_blocks = section.get("content", [])

                if heading:
                    heading_text = heading + "\n"
                    requests.append(
                        {"insertText": {"location": {"index": cursor}, "text": heading_text}}
                    )
                    h_start = cursor
                    h_end = cursor + len(heading_text)
                    requests.append(
                        {
                            "updateParagraphStyle": {
                                "range": {"startIndex": h_start, "endIndex": h_end},
                                "paragraphStyle": {
                                    "namedStyleType": _heading_style(level),
                                    "spaceAbove": {"magnitude": 12, "unit": "PT"},
                                    "spaceBelow": {"magnitude": 4, "unit": "PT"},
                                },
                                "fields": "namedStyleType,spaceAbove,spaceBelow",
                            }
                        }
                    )
                    cursor = h_end

                for block in content_blocks:
                    block_type = block.get("type", "paragraph")

                    if block_type == "paragraph":
                        text = block.get("text", "")
                        if not text:
                            continue
                        para_text = text + "\n"
                        plain, formatting_runs = _parse_inline_formatting(para_text)
                        requests.append(
                            {"insertText": {"location": {"index": cursor}, "text": plain}}
                        )
                        p_start = cursor
                        p_end = cursor + len(plain)
                        requests.append(
                            {
                                "updateParagraphStyle": {
                                    "range": {"startIndex": p_start, "endIndex": p_end},
                                    "paragraphStyle": {
                                        "namedStyleType": "NORMAL_TEXT",
                                        "spaceBelow": {"magnitude": 6, "unit": "PT"},
                                    },
                                    "fields": "namedStyleType,spaceBelow",
                                }
                            }
                        )
                        for run_start, run_end, style in formatting_runs:
                            abs_start = cursor + run_start
                            abs_end = cursor + run_end
                            if style == "bold":
                                requests.append(
                                    {
                                        "updateTextStyle": {
                                            "range": {
                                                "startIndex": abs_start,
                                                "endIndex": abs_end,
                                            },
                                            "textStyle": {"bold": True},
                                            "fields": "bold",
                                        }
                                    }
                                )
                            elif style == "italic":
                                requests.append(
                                    {
                                        "updateTextStyle": {
                                            "range": {
                                                "startIndex": abs_start,
                                                "endIndex": abs_end,
                                            },
                                            "textStyle": {"italic": True},
                                            "fields": "italic",
                                        }
                                    }
                                )
                        cursor = p_end

                    elif block_type in ("bullets", "numbered_list"):
                        items = block.get("items", [])
                        if not items:
                            continue
                        list_start = cursor
                        for item in items:
                            item_text = str(item) + "\n"
                            requests.append(
                                {
                                    "insertText": {
                                        "location": {"index": cursor},
                                        "text": item_text,
                                    }
                                }
                            )
                            cursor += len(item_text)
                        list_end = cursor
                        preset = (
                            "BULLET_DISC_CIRCLE_SQUARE"
                            if block_type == "bullets"
                            else "NUMBERED_DECIMAL_NESTED"
                        )
                        requests.append(
                            {
                                "createParagraphBullets": {
                                    "range": {
                                        "startIndex": list_start,
                                        "endIndex": list_end,
                                    },
                                    "bulletPreset": preset,
                                }
                            }
                        )

                    elif block_type == "table":
                        t_headers = block.get("headers", [])
                        t_rows = block.get("rows", [])
                        if not t_headers:
                            continue
                        num_cols = len(t_headers)
                        num_rows = 1 + len(t_rows)
                        requests.append(
                            {
                                "insertTable": {
                                    "rows": num_rows,
                                    "columns": num_cols,
                                    "location": {"index": cursor},
                                }
                            }
                        )
                        if requests:
                            _rate_limited_batch(requests)
                            requests = []

                        doc_state = self._docs_service.documents().get(documentId=doc_id).execute()
                        table_element = None
                        for elem in doc_state.get("body", {}).get("content", []):
                            if "table" in elem:
                                table_element = elem
                        if table_element:
                            table = table_element["table"]
                            all_table_rows = [t_headers] + t_rows
                            cell_inserts: List[tuple] = []
                            for ri, row_data in enumerate(all_table_rows):
                                if ri >= len(table.get("tableRows", [])):
                                    break
                                table_row = table["tableRows"][ri]
                                for ci, cell_val in enumerate(row_data):
                                    if ci >= len(table_row.get("tableCells", [])):
                                        break
                                    cell = table_row["tableCells"][ci]
                                    cell_content = cell.get("content", [])
                                    if cell_content:
                                        cell_idx = cell_content[0].get("startIndex", 0)
                                        cell_text = str(cell_val)
                                        if cell_text:
                                            cell_inserts.append((cell_idx, cell_text))
                            cell_inserts.sort(key=lambda x: x[0], reverse=True)
                            for cell_idx, cell_text in cell_inserts:
                                requests.append(
                                    {
                                        "insertText": {
                                            "location": {"index": cell_idx},
                                            "text": cell_text,
                                        }
                                    }
                                )
                            if requests:
                                _rate_limited_batch(requests)
                                requests = []

                            doc_state2 = (
                                self._docs_service.documents().get(documentId=doc_id).execute()
                            )
                            table_elem2 = None
                            for elem in doc_state2.get("body", {}).get("content", []):
                                if "table" in elem:
                                    table_elem2 = elem
                            if table_elem2 and table_elem2["table"].get("tableRows"):
                                header_row = table_elem2["table"]["tableRows"][0]
                                for hcell in header_row.get("tableCells", []):
                                    for para in hcell.get("content", []):
                                        si = para.get("startIndex", 0)
                                        ei = para.get("endIndex", si)
                                        if ei > si:
                                            requests.append(
                                                {
                                                    "updateTextStyle": {
                                                        "range": {
                                                            "startIndex": si,
                                                            "endIndex": ei,
                                                        },
                                                        "textStyle": {"bold": True},
                                                        "fields": "bold",
                                                    }
                                                }
                                            )
                            if requests:
                                _rate_limited_batch(requests)
                                requests = []

                        doc_state = self._docs_service.documents().get(documentId=doc_id).execute()
                        body_content = doc_state.get("body", {}).get("content", [])
                        if body_content:
                            last_elem = body_content[-1]
                            cursor = last_elem.get("endIndex", cursor) - 1
                        requests.append(
                            {"insertText": {"location": {"index": cursor}, "text": "\n"}}
                        )
                        cursor += 1

            if requests:
                _rate_limited_batch(requests)

            # Insert table of contents if requested
            if include_toc:
                self._insert_table_of_contents(doc_id, batch_fn=_rate_limited_batch)

            docs_url = f"https://docs.google.com/document/d/{doc_id}/edit"
            logger.info(
                "google_doc_updated",
                doc_id=doc_id,
                sections=len(sections),
                workspace=workspace_id,
            )
            return {
                "status": "updated",
                "google_docs_url": docs_url,
                "google_docs_id": doc_id,
            }

        except Exception as e:
            logger.error("google_doc_update_failed", error=str(e))
            return {"status": "error", "message": str(e)}


# ══════════════════════════════════════════════════════════════════════════
# Helpers
# ══════════════════════════════════════════════════════════════════════════


def _heading_style(level: int) -> str:
    """Map heading level (1-6) to a Google Docs named style."""
    mapping = {
        1: "HEADING_1",
        2: "HEADING_2",
        3: "HEADING_3",
        4: "HEADING_4",
        5: "HEADING_5",
        6: "HEADING_6",
    }
    return mapping.get(level, "HEADING_1")


def _normalize_sections(structure: Dict[str, Any], skip_toc: bool = False) -> List[Dict[str, Any]]:
    """Normalize sections to use the rich ``content`` array format.

    Converts old-style ``{"body": "..."}`` sections to
    ``{"content": [{"type": "paragraph", "text": "..."}]}``.

    When *skip_toc* is True, sections with heading "Table of Contents"
    or content type ``"table_of_contents"`` are filtered out (the TOC
    will be auto-generated after all content is placed).
    """
    sections = structure.get("sections", [])
    normalized: List[Dict[str, Any]] = []
    for section in sections:
        # Skip manual TOC sections when auto-TOC is enabled
        if skip_toc:
            heading = section.get("heading", "").strip().lower()
            if heading == "table of contents":
                continue
            content_blocks = section.get("content", [])
            if len(content_blocks) == 1 and content_blocks[0].get("type") == "table_of_contents":
                continue
        s = dict(section)
        if "content" not in s:
            body = s.pop("body", "")
            if body:
                s["content"] = [{"type": "paragraph", "text": body}]
            else:
                s["content"] = []
        normalized.append(s)
    return normalized


def _parse_inline_formatting(text: str) -> tuple:
    """Parse ``**bold**`` and ``*italic*`` markers from *text*.

    Returns ``(plain_text, [(start, end, style), ...])``.
    Runs are non-overlapping and refer to positions in *plain_text*.
    """
    runs: List[tuple] = []
    result_chars: List[str] = []
    i = 0
    length = len(text)
    while i < length:
        # Bold: **...**
        if i + 1 < length and text[i] == "*" and text[i + 1] == "*":
            end_marker = text.find("**", i + 2)
            if end_marker != -1:
                start_pos = len(result_chars)
                inner = text[i + 2 : end_marker]
                result_chars.extend(inner)
                runs.append((start_pos, start_pos + len(inner), "bold"))
                i = end_marker + 2
                continue
        # Italic: *...*  (single asterisk, not followed by another)
        if text[i] == "*" and (i + 1 >= length or text[i + 1] != "*"):
            end_marker = text.find("*", i + 1)
            if end_marker != -1 and (end_marker + 1 >= length or text[end_marker + 1] != "*"):
                start_pos = len(result_chars)
                inner = text[i + 1 : end_marker]
                result_chars.extend(inner)
                runs.append((start_pos, start_pos + len(inner), "italic"))
                i = end_marker + 1
                continue
        result_chars.append(text[i])
        i += 1
    return ("".join(result_chars), runs)


def _doc_structure_format_instructions() -> str:
    return """Return ONLY a valid JSON object with this structure:
{
  "title": "Document Title",
  "sections": [
    {
      "heading": "Section Heading",
      "level": 1,
      "content": [
        {"type": "paragraph", "text": "Use **bold** for emphasis and *italic* for terms."},
        {"type": "bullets", "items": ["First point", "Second point", "Third point"]},
        {"type": "numbered_list", "items": ["Step 1", "Step 2", "Step 3"]},
        {"type": "table", "headers": ["Column A", "Column B"], "rows": [["val1", "val2"]]}
      ]
    },
    {
      "heading": "Subsection",
      "level": 2,
      "content": [
        {"type": "paragraph", "text": "More content here with **bold** highlights."}
      ]
    }
  ]
}

Use level 1 for major sections, level 2 for subsections.
Use **bold** for emphasis and *italic* for key terms within paragraph text.
Use "bullets" for unordered lists, "numbered_list" for ordered steps.
Use "table" for structured data with headers and rows.
Write substantive content — at least 2-3 sentences per paragraph. Be specific and actionable.

ABBREVIATIONS: Always spell out abbreviations/acronyms on first use with the abbreviation in
parentheses — e.g., "Total Addressable Market (TAM)". After first use, use the abbreviation only.

TABLE OF CONTENTS: Set "include_toc": true at the top level (next to "title") to auto-generate
a linked table of contents from all headings. Do NOT manually create a TOC section — it will
be generated automatically with clickable links to each heading. Note: the Google Docs API
cannot insert a native TOC with page numbers (known Google limitation). The generated TOC
includes an instruction for the user to replace it with a native TOC if page numbers are needed.

BACKWARD COMPATIBILITY: You may also use the simpler format with "body" string instead of
"content" array, but the rich format above is strongly preferred."""


def _parse_doc_structure(raw: str) -> Dict[str, Any]:
    """Parse LLM output into a document structure dict."""
    # Try direct JSON parse
    try:
        return json.loads(raw)
    except (json.JSONDecodeError, TypeError):
        pass

    # Try extracting JSON from markdown code block
    match = re.search(r"```(?:json)?\s*\n?(.*?)\n?```", raw, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except (json.JSONDecodeError, TypeError):
            pass

    # Try stripping garbage prefix (e.g., "Here is the JSON:\n{...}")
    json_start = raw.find("{")
    if json_start > 0:
        try:
            return json.loads(raw[json_start:])
        except (json.JSONDecodeError, TypeError):
            pass

    # Fallback: plain text — skip lines that look like raw JSON fragments
    logger.warning("doc_structure_parse_fallback", raw_length=len(raw))
    _JSON_LINE = re.compile(r"^\s*[\[\]{},]+\s*$")
    lines = [
        line.strip() for line in raw.split("\n") if line.strip() and not _JSON_LINE.match(line)
    ]
    return {
        "title": lines[0] if lines else "Untitled Document",
        "sections": [
            {
                "heading": "Content",
                "body": "\n".join(lines[1:]) if len(lines) > 1 else "Content to be added.",
                "level": 1,
            }
        ],
    }
